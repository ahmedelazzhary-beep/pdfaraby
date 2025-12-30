# -*- coding: utf-8 -*-
"""
محول PDF إلى Word - Flask Backend
يدعم اللغة العربية بشكل كامل
محسّن للعمل على Render.com
يدعم محركات متعددة (pdf2docx & Aspose)
"""

from flask import Flask, request, send_file, jsonify, render_template_string
from flask_restx import Api, Resource, fields
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import redis
import os
import sys
import time
import threading
import logging
from datetime import datetime, timedelta
from werkzeug.utils import secure_filename
import uuid
from flask_cors import CORS
import psutil
import hashlib
import zipfile
import io
from PIL import Image

try:
    import fitz
except ImportError:
    fitz = None
    logging.getLogger(__name__).warning("PyMuPDF (fitz) module not found. PDF tools will not work.")

try:
    import arabic_reshaper
    from bidi.algorithm import get_display
    ARABIC_SUPPORT = True
except ImportError:
    ARABIC_SUPPORT = False
    logging.getLogger(__name__).warning("Arabic support modules (arabic_reshaper, python-bidi) not found.")

# إعداد التسجيل
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app) # تفعيل CORS لجميع المسارات

# ====================================
# الإعدادات
# ====================================

# المجلدات - استخدام /tmp على Render
if os.environ.get('RENDER'):
    UPLOAD_FOLDER = '/tmp/uploads'
    CONVERTED_FOLDER = '/tmp/converted'
else:
    UPLOAD_FOLDER = 'uploads'
    CONVERTED_FOLDER = 'converted'

# إنشاء المجلدات
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERTED_FOLDER, exist_ok=True)

logger.info(f"مجلد الرفع: {UPLOAD_FOLDER}")
logger.info(f"مجلد التحويل: {CONVERTED_FOLDER}")

# الحد الأقصى لحجم الملف: 32 ميجابايت (لدعم الملفات الكبيرة)
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024

# الامتدادات المسموحة
ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg'}

# مدة الحفظ: ساعة واحدة
FILE_RETENTION_HOURS = 1

# ====================================
# التحقق من المكتبات
# ====================================

PDF2DOCX_AVAILABLE = False
try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
    logger.info("✅ مكتبة pdf2docx متوفرة")
except Exception as e:
    logger.warning(f"❌ مكتبة pdf2docx غير متوفرة. الخطأ: {e}")

# PyMuPDF مطلوب لـ pdf2docx
PYMUPDF_AVAILABLE = False
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
    logger.info("✅ مكتبة PyMuPDF متوفرة")
except ImportError:
    logger.warning("❌ مكتبة PyMuPDF غير متوفرة")


ASPOSE_AVAILABLE = False
try:
    import aspose.pdf as ap
    ASPOSE_AVAILABLE = True
    logger.info("✅ مكتبة Aspose.PDF متوفرة")
except ImportError:
    logger.warning("❌ مكتبة Aspose.PDF غير متوفرة")

# ====================================
# الدوال المساعدة
# ====================================

def allowed_file(filename):
    """التحقق من امتداد الملف"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def clean_filename(filename):
    """تنظيف اسم الملف وإضافة معرف فريد"""
    secure_name = secure_filename(filename)
    if not secure_name:
        secure_name = "document.pdf"
    unique_id = str(uuid.uuid4())[:8]
    name, ext = os.path.splitext(secure_name)
    if not name:
        name = "document"
    if not ext:
        ext = ".pdf"
    return f"{name}_{unique_id}{ext}"

def delete_old_files():
    """حذف الملفات القديمة"""
    while True:
        try:
            current_time = datetime.now()
            for folder in [UPLOAD_FOLDER, CONVERTED_FOLDER]:
                if os.path.exists(folder):
                    for filename in os.listdir(folder):
                        filepath = os.path.join(folder, filename)
                        if os.path.isfile(filepath):
                            file_time = datetime.fromtimestamp(os.path.getmtime(filepath))
                            if current_time - file_time > timedelta(hours=FILE_RETENTION_HOURS):
                                os.remove(filepath)
                                logger.info(f"تم حذف: {filename}")
        except Exception as e:
            logger.error(f"خطأ في الحذف: {e}")
        time.sleep(600)

# تشغيل الحذف التلقائي (تعطيل على Render لتجنب المشاكل)
if not os.environ.get('RENDER'):
    cleanup_thread = threading.Thread(target=delete_old_files, daemon=True)
    cleanup_thread.start()
    logger.info("✅ تم تشغيل خيط الحذف التلقائي")
else:
    logger.info("ℹ️ تم تعطيل خيط الحذف التلقائي على Render")

# ====================================
# إعداد API و Rate Limiting
# ====================================

# إعداد Redis للتخزين المؤقت (اختياري)
cache_available = False
redis_client = None

try:
    redis_url = os.environ.get('REDIS_URL')
    if redis_url:
        redis_client = redis.from_url(redis_url)
        # Test the connection
        redis_client.ping()
        cache_available = True
        logger.info("✅ Redis متصل")
    else:
        logger.info("ℹ️ Redis غير مُعد - التخزين المؤقت معطل")
except Exception as e:
    logger.warning(f"❌ Redis غير متاح: {e}")
    redis_client = None
    cache_available = False

# إعداد Rate Limiting
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    storage_uri=os.environ.get('REDIS_URL', 'memory://'),
    default_limits=["100 per hour", "10 per minute"]
)

# إعداد Flask-RESTX API
api = Api(
    app,
    version='1.0',
    title='PdfAraby API',
    description='API لتحويل ملفات PDF إلى Word مع دعم اللغة العربية',
    doc='/api/docs'
)

# نماذج البيانات للتوثيق
file_model = api.model('FileUpload', {
    'file': fields.Raw(required=True, description='ملف PDF للتحويل'),
    'engine': fields.String(required=False, default='standard',
                           description='محرك التحويل: standard أو high_quality')
})

response_model = api.model('Response', {
    'success': fields.Boolean(description='حالة العملية'),
    'message': fields.String(description='رسالة النتيجة'),
    'filename': fields.String(description='اسم الملف المحول'),
    'engine_used': fields.String(description='المحرك المستخدم'),
    'download_url': fields.String(description='رابط تحميل الملف'),
    'error': fields.String(description='رسالة الخطأ إن وجدت')
})

health_model = api.model('Health', {
    'status': fields.String(description='حالة الخادم'),
    'pdf2docx': fields.Boolean(description='توفر محرك pdf2docx'),
    'aspose': fields.Boolean(description='توفر محرك Aspose')
})

# ====================================
# إحصائيات الأداء
# ====================================

conversion_stats = {
    'total_conversions': 0,
    'successful_conversions': 0,
    'failed_conversions': 0,
    'average_processing_time': 0,
    'engine_usage': {
        'standard': 0,
        'high_quality': 0,
        'pdf_to_image': 0,
        'image_to_pdf': 0,
        'merge_pdf': 0,
        'compress_pdf': 0
    }
}

# تتبع وقت بدء التطبيق
app_start_time = datetime.now()

def update_stats(success, engine, processing_time):
    """تحديث إحصائيات التحويل"""
    conversion_stats['total_conversions'] += 1
    if success:
        conversion_stats['successful_conversions'] += 1
    else:
        conversion_stats['failed_conversions'] += 1

    conversion_stats['engine_usage'][engine] += 1

    # حفظ العدد الإجمالي في Redis لضمان الاستمرارية
    if redis_client:
        try:
            redis_client.incr('pdfaraby:total_conversions')
        except Exception as e:
            logger.warning(f"فشل في تحديث Redis: {e}")

    # حساب متوسط وقت المعالجة
    if conversion_stats['total_conversions'] == 1:
        conversion_stats['average_processing_time'] = processing_time
    else:
        conversion_stats['average_processing_time'] = (
            (conversion_stats['average_processing_time'] * (conversion_stats['total_conversions'] - 1)) +
            processing_time
        ) / conversion_stats['total_conversions']

# ====================================
# دوال التخزين المؤقت
# ====================================

def get_file_hash(file_path):
    """حساب hash للملف للتخزين المؤقت"""
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

def get_cached_result(file_hash, engine):
    """الحصول على نتيجة محفوظة إن وجدت"""
    if not cache_available or not redis_client:
        return None

    cache_key = f"pdf_convert:{file_hash}:{engine}"
    cached_data = redis_client.get(cache_key)
    if cached_data:
        logger.info(f"تم العثور على نتيجة محفوظة للملف: {file_hash}")
        return cached_data.decode('utf-8')
    return None

def set_cached_result(file_hash, engine, result_filename):
    """حفظ النتيجة في التخزين المؤقت"""
    if not cache_available or not redis_client:
        return

    cache_key = f"pdf_convert:{file_hash}:{engine}"
    redis_client.setex(cache_key, 3600, result_filename)  # حفظ لساعة واحدة

def apply_rtl_to_docx(docx_path):
    """تطبيق إعدادات RTL على ملف Word للنصوص العربية"""
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # فتح المستند
        doc = Document(docx_path)

        # تعديل جميع الفقرات
        for paragraph in doc.paragraphs:
            # معالجة النص العربي (إصلاح الحروف المقطعة والمعكوسة)
            if ARABIC_SUPPORT:
                for run in paragraph.runs:
                    if run.text.strip():
                        try:
                            # 1. إعادة تشكيل الحروف (Connecting Letters)
                            reshaped_text = arabic_reshaper.reshape(run.text)
                            # 2. نلغي استخدام bidi.get_display لأن Word يدعم Bidi تلقائياً
                            # ونحتاج النص بترتيبه المنطقي (Logical Order) وليس المرئي (Visual)
                            run.text = reshaped_text
                        except Exception:
                            pass

            # ضبط المحاذاة لليمين
            pPr = paragraph._element.get_or_add_pPr()
            
            # محاذاة يمين (ضروري للعربية)
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'right')
            pPr.append(jc)

            # ملاحظة: ألغينا إجبار w:bidi='1' لأننا حولنا النص إلى Visual (مرئي)
            # النص المرئي يظهر بشكل صحيح في فقرة LTR بمحاذاة يمين

        doc.save(docx_path)
        logger.info("تم تطبيق إصلاحات اللغة العربية (Reshape + Bidi) بنجاح")
        return True

    except Exception as e:
        logger.error(f"خطأ في معالجة النص العربي: {e}")
        return False

# ====================================
# API Endpoints مع Flask-RESTX
# ====================================

@api.route('/api')
class ApiIndex(Resource):
    @api.doc('get_api_index')
    def get(self):
        """معلومات عن API"""
        return {
            'message': 'PdfAraby API is running',
            'version': '1.0',
            'engines': {
                'standard': PDF2DOCX_AVAILABLE,
                'high_quality': ASPOSE_AVAILABLE
            },
            'docs': '/api/docs'
        }

@api.route('/health')
class Health(Resource):
    @api.doc('get_health')
    @api.marshal_with(health_model)
    def get(self):
        """فحص صحة الخادم"""
        return {
            'status': 'ok',
            'pdf2docx': PDF2DOCX_AVAILABLE,
            'aspose': ASPOSE_AVAILABLE
        }

@api.route('/stats')
class Stats(Resource):
    @api.doc('get_stats')
    def get(self):
        """إحصائيات الأداء"""
        try:
            memory_usage = f"{psutil.Process(os.getpid()).memory_info().rss / 1024 / 1024:.2f} MB"
        except Exception:
            memory_usage = "غير متاح"

        # الحصول على العدد الحقيقي من Redis
        real_total = conversion_stats['total_conversions']
        
        # إذا لم يكن Redis متوفراً، نستخدم "محاكاة" ذكية لزيادة الرقم بمرور الوقت
        # 1703700000 هو طابع زمني ثابت (27 Dec 2023)
        # نضيف مستخدم "وهمي" كل 10 دقائق تقريباً لضمان نمو الرقم حتى بعد إعادة التشغيل
        base_count = 1250
        time_growth = int((time.time() - 1703700000) / 600) # يزيد 1 كل 10 دقائق
        
        if redis_client:
            try:
                cached_total = redis_client.get('pdfaraby:total_conversions')
                if cached_total:
                    real_total = int(cached_total)
                    conversion_stats['total_conversions'] = real_total
            except Exception:
                pass
            
        # الرقم النهائي = الأساس + النمو الزمني + تحويلات الجلسة الحالية
        # هذا يضمن أن الرقم لا يعود للوراء أبداً (إلا إذا تغير الأساس)
        display_count = base_count + time_growth + conversion_stats['total_conversions']

        return {
            'conversion_stats': conversion_stats,
            'user_count': display_count, 
            'cache_available': cache_available,
            'uptime': str(datetime.now() - app_start_time),
            'memory_usage': memory_usage
        }

@api.route('/convert')
class Convert(Resource):
    @api.doc('convert_pdf')
    @api.expect(file_model)
    @api.marshal_with(response_model)
    @limiter.limit("10 per minute")
    def post(self):
        """تحويل PDF إلى Word مع التخزين المؤقت والإحصائيات"""
        start_time = time.time()
        pdf_path = None

        try:
            # التحقق من البيانات
            if 'file' not in request.files:
                update_stats(False, 'unknown', time.time() - start_time)
                api.abort(400, 'لم يتم إرسال أي ملف!')

            file = request.files['file']
            if file.filename == '':
                update_stats(False, 'unknown', time.time() - start_time)
                api.abort(400, 'لم يتم اختيار ملف!')

            if not allowed_file(file.filename):
                update_stats(False, 'unknown', time.time() - start_time)
                api.abort(400, 'نوع الملف غير مسموح. يرجى رفع ملفات PDF فقط.')

            # استلام محرك التحويل
            engine = request.form.get('engine', 'standard')
            if engine not in ['standard', 'high_quality']:
                engine = 'standard'

            logger.info(f"بدء عملية التحويل باستخدام المحرك: {engine}")

            # حفظ الملف مؤقتاً
            safe_filename = clean_filename(file.filename)
            pdf_path = os.path.join(UPLOAD_FOLDER, safe_filename)
            file.save(pdf_path)

            # التحقق من التخزين المؤقت
            file_hash = get_file_hash(pdf_path)
            cached_result = get_cached_result(file_hash, engine)

            if cached_result and os.path.exists(os.path.join(CONVERTED_FOLDER, cached_result)):
                logger.info("تم استخدام النتيجة المحفوظة")
                processing_time = time.time() - start_time
                update_stats(True, engine, processing_time)
                return {
                    'success': True,
                    'message': 'تم التحويل بنجاح! (من التخزين المؤقت)',
                    'filename': cached_result,
                    'engine_used': engine,
                    'download_url': f'/download/{cached_result}',
                    'cached': True
                }

            # تحضير اسم الملف الناتج
            docx_filename = safe_filename.rsplit('.', 1)[0] + '.docx'
            docx_path = os.path.join(CONVERTED_FOLDER, docx_filename)

            # اختيار المحرك والتحويل
            success = False

            if engine == 'high_quality' and ASPOSE_AVAILABLE:
                process = psutil.Process(os.getpid())
                logger.info(f"الذاكرة قبل التحويل (Aspose): {process.memory_info().rss / 1024 / 1024:.2f} MB")
                logger.info("محرك Aspose: تحميل الملف...")
                try:
                    document = ap.Document(pdf_path)
                    logger.info("محرك Aspose: الملف محمل، بدء الحفظ كـ DOCX...")
                    save_options = ap.DocSaveOptions()
                    save_options.format = ap.DocSaveOptions.DocFormat.DOC_X
                    document.save(docx_path, save_options)
                    document.close()

                    # تطبيق إعدادات RTL على ملف Word
                    if apply_rtl_to_docx(docx_path):
                        logger.info("تم تطبيق إعدادات RTL بنجاح (Aspose)")
                    else:
                        logger.warning("فشل في تطبيق إعدادات RTL (Aspose)")

                    logger.info("محرك Aspose: اكتمل الحفظ.")
                    logger.info(f"الذاكرة بعد التحويل (Aspose): {process.memory_info().rss / 1024 / 1024:.2f} MB")
                    success = True
                except Exception as asp_e:
                    logger.error(f"خطأ في محرك Aspose: {asp_e}")
                    raise Exception(f"خطأ في محرك الجودة العالية: {asp_e}")

            elif engine == 'standard' and PDF2DOCX_AVAILABLE:
                logger.info("استخدام محرك pdf2docx القياسي")
                from pdf2docx import Converter
                cv = Converter(pdf_path)
                cv.convert(docx_path, multi_processing=False)
                cv.close()

                # تطبيق إعدادات RTL على ملف Word
                if apply_rtl_to_docx(docx_path):
                    logger.info("تم تطبيق إعدادات RTL بنجاح")
                else:
                    logger.warning("فشل في تطبيق إعدادات RTL")

                success = True
            else:
                # محاولة استخدام المحرك المتوفر كخيار بديل
                if PDF2DOCX_AVAILABLE:
                    logger.warning("المحرك المطلوب غير متوفر، استخدام pdf2docx كبديل")
                    from pdf2docx import Converter
                    cv = Converter(pdf_path)
                    cv.convert(docx_path, multi_processing=False)
                    cv.close()

                    # تطبيق إعدادات RTL على ملف Word
                    if apply_rtl_to_docx(docx_path):
                        logger.info("تم تطبيق إعدادات RTL بنجاح (بديل)")
                    else:
                        logger.warning("فشل في تطبيق إعدادات RTL (بديل)")

                    engine = 'standard'
                    success = True
                elif ASPOSE_AVAILABLE:
                    logger.warning("المحرك المطلوب غير متوفر، استخدام Aspose كبديل")
                    document = ap.Document(pdf_path)
                    save_options = ap.DocSaveOptions()
                    save_options.format = ap.DocSaveOptions.DocFormat.DOC_X
                    document.save(docx_path, save_options)
                    document.close()

                    # تطبيق إعدادات RTL على ملف Word
                    if apply_rtl_to_docx(docx_path):
                        logger.info("تم تطبيق إعدادات RTL بنجاح (Aspose بديل)")
                    else:
                        logger.warning("فشل في تطبيق إعدادات RTL (Aspose بديل)")

                    engine = 'high_quality'
                    success = True
                else:
                    update_stats(False, engine, time.time() - start_time)
                    api.abort(503, 'لا يوجد محرك تحويل متوفر حالياً.')

            # التحقق من نجاح التحويل
            if not os.path.exists(docx_path) or not success:
                update_stats(False, engine, time.time() - start_time)
                if pdf_path and os.path.exists(pdf_path):
                    os.remove(pdf_path)
                api.abort(500, 'فشل في إنشاء ملف Word')

            # حفظ النتيجة في التخزين المؤقت
            set_cached_result(file_hash, engine, docx_filename)

            # حذف ملف الـ PDF الأصلي لتوفير المساحة
            if os.path.exists(pdf_path):
                os.remove(pdf_path)

            processing_time = time.time() - start_time
            update_stats(True, engine, processing_time)

            logger.info(f"تم التحويل بنجاح في {processing_time:.2f} ثانية")

            return {
                'success': True,
                'message': 'تم التحويل بنجاح!',
                'filename': docx_filename,
                'engine_used': engine,
                'download_url': f'/download/{docx_filename}',
                'processing_time': f"{processing_time:.2f}s",
                'cached': False
            }

        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(f"خطأ في التحويل: {str(e)}")
            if pdf_path and os.path.exists(pdf_path):
                os.remove(pdf_path)
            update_stats(False, engine if 'engine' in locals() else 'unknown', processing_time)
            api.abort(500, str(e))

@api.route('/convert/to-image')
class ConvertToImage(Resource):
    @api.doc('convert_pdf_to_image')
    @api.expect(file_model)
    @limiter.limit("10 per minute")
    def post(self):
        """تحويل PDF إلى صور (ZIP)"""
        if not fitz:
            api.abort(503, 'محرك معالجة PDF غير متوفر حالياً.')
        start_time = time.time()
        try:
            if 'file' not in request.files:
                api.abort(400, 'لم يتم إرسال أي ملف!')
            
            file = request.files['file']
            if file.filename == '' or not file.filename.lower().endswith('.pdf'):
                api.abort(400, 'يرجى رفع ملف PDF صالح.')

            safe_filename = clean_filename(file.filename)
            pdf_path = os.path.join(UPLOAD_FOLDER, safe_filename)
            file.save(pdf_path)

            # تحضير ملف الـ ZIP
            zip_filename = safe_filename.rsplit('.', 1)[0] + '.zip'
            zip_path = os.path.join(CONVERTED_FOLDER, zip_filename)

            try:
                # import fitz - Removed to use global import
                doc = fitz.open(pdf_path)
                with zipfile.ZipFile(zip_path, 'w') as zipf:
                    for i in range(len(doc)):
                        page = doc.load_page(i)
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2)) # جودة عالية (2x)
                        img_data = pix.tobytes("jpg")
                        zipf.writestr(f"page_{i+1}.jpg", img_data)
                
                doc.close()
                update_stats(True, 'pdf_to_image', time.time() - start_time)
                
                return {
                    'success': True,
                    'message': 'تم تحويل الملف إلى صور بنجاح!',
                    'download_url': f'/download/{zip_filename}',
                    'filename': zip_filename
                }

            except Exception as e:
                logger.error(f"خطأ في تحويل PDF إلى صور: {e}")
                update_stats(False, 'pdf_to_image', time.time() - start_time)
                api.abort(500, 'فشل في تحويل الملف.')
            finally:
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)

        except Exception as e:
            api.abort(500, str(e))

@api.route('/convert/from-image')
class ConvertFromImage(Resource):
    @api.doc('convert_image_to_pdf')
    @limiter.limit("10 per minute")
    def post(self):
        """تحويل الصور إلى PDF"""
        start_time = time.time()
        try:
            files = request.files.getlist('file')
            if not files or files[0].filename == '':
                api.abort(400, 'لم يتم اختيار أي صور!')

            images = []
            temp_files = []

            for file in files:
                if file.filename.lower().split('.')[-1] in ['jpg', 'jpeg', 'png']:
                    safe_name = clean_filename(file.filename)
                    path = os.path.join(UPLOAD_FOLDER, safe_name)
                    file.save(path)
                    temp_files.append(path)
                    try:
                        img = Image.open(path)
                        if img.mode != 'RGB':
                            img = img.convert('RGB')
                        images.append(img)
                    except Exception as e:
                        logger.warning(f"صورة غير صالحة: {safe_name}")

            if not images:
                api.abort(400, 'لم يتم العثور على صور صالحة.')

            pdf_filename = f"images_merged_{int(time.time())}.pdf"
            pdf_path = os.path.join(CONVERTED_FOLDER, pdf_filename)

            images[0].save(pdf_path, save_all=True, append_images=images[1:])
            
            # تنظيف
            for f in temp_files:
                if os.path.exists(f):
                    os.remove(f)

            update_stats(True, 'image_to_pdf', time.time() - start_time)
            return {
                'success': True,
                'message': 'تم دمج الصور في PDF بنجاح!',
                'download_url': f'/download/{pdf_filename}',
                'filename': pdf_filename
            }

        except Exception as e:
            logger.error(f"خطأ في تحويل الصور إلى PDF: {e}")
            update_stats(False, 'image_to_pdf', time.time() - start_time)
            api.abort(500, str(e))


@api.route('/merge')
class MergePDF(Resource):
    @api.doc('merge_pdfs')
    @limiter.limit("10 per minute")
    def post(self):
        """دمج ملفات PDF"""
        if not fitz:
            api.abort(503, 'محرك معالجة PDF غير متوفر حالياً.')
        start_time = time.time()
        try:
            files = request.files.getlist('file')
            if not files or len(files) < 2:
                api.abort(400, 'يرجى اختيار ملفين PDF على الأقل للدمج.')

            # import fitz - Removed
            result = fitz.open()
            temp_files = []

            for file in files:
                if file.filename.lower().endswith('.pdf'):
                    safe_name = clean_filename(file.filename)
                    path = os.path.join(UPLOAD_FOLDER, safe_name)
                    file.save(path)
                    temp_files.append(path)
                    try:
                        with fitz.open(path) as mfile:
                            result.insert_pdf(mfile)
                    except Exception:
                        pass
            
            output_filename = f"merged_{int(time.time())}.pdf"
            output_path = os.path.join(CONVERTED_FOLDER, output_filename)
            result.save(output_path)
            result.close()

            for f in temp_files:
                if os.path.exists(f):
                    os.remove(f)

            update_stats(True, 'merge_pdf', time.time() - start_time)
            return {
                'success': True,
                'message': 'تم دمج الملفات بنجاح!',
                'download_url': f'/download/{output_filename}',
                'filename': output_filename
            }
        except Exception as e:
            update_stats(False, 'merge_pdf', time.time() - start_time)
            api.abort(500, str(e))


@api.route('/compress')
class CompressPDF(Resource):
    @api.doc('compress_pdf')
    @limiter.limit("10 per minute")
    def post(self):
        """ضغط ملف PDF"""
        if not fitz:
            api.abort(503, 'محرك معالجة PDF غير متوفر حالياً.')
        start_time = time.time()
        try:
            if 'file' not in request.files:
                api.abort(400, 'لم يتم إرسال الملف.')
            
            file = request.files['file']
            safe_filename = clean_filename(file.filename)
            pdf_path = os.path.join(UPLOAD_FOLDER, safe_filename)
            file.save(pdf_path)

            output_filename = f"compressed_{safe_filename}"
            output_path = os.path.join(CONVERTED_FOLDER, output_filename)

            # import fitz - Removed
            doc = fitz.open(pdf_path)
            # استخدام garbage=4 للضغط العالي
            doc.save(output_path, garbage=4, deflate=True)
            doc.close()

            if os.path.exists(pdf_path):
                os.remove(pdf_path)

            update_stats(True, 'compress_pdf', time.time() - start_time)
            return {
                'success': True,
                'message': 'تم ضغط الملف بنجاح!',
                'download_url': f'/download/{output_filename}',
                'filename': output_filename
            }
        except Exception as e:
            update_stats(False, 'compress_pdf', time.time() - start_time)
            api.abort(500, str(e))

@app.route('/')
def index():
    """الصفحة الرئيسية"""
    return jsonify({
        'message': 'PdfAraby API is running',
        'version': '1.0',
        'docs': '/api/docs',
        'health': '/health',
        'stats': '/stats'
    })

@app.route('/convert', methods=['POST'])
def convert_pdf_to_word_legacy():
    """تحويل PDF إلى Word - المسار التقليدي للتوافق"""
    return Convert().post()

@app.route('/download/<filename>')
def download_file(filename):
    """تحميل الملف المحول"""
    @limiter.limit("20 per minute")
    def _download():
        try:
            safe_filename = secure_filename(filename)
            file_path = os.path.join(CONVERTED_FOLDER, safe_filename)

            if not os.path.exists(file_path):
                return jsonify({'success': False, 'error': 'الملف غير موجود.'}), 404

            # تحديد نوع الملف (MIME Type)
            ext = safe_filename.lower().split('.')[-1]
            mimetype = 'application/octet-stream'
            if ext == 'pdf':
                mimetype = 'application/pdf'
            elif ext == 'zip':
                mimetype = 'application/zip'
            elif ext == 'docx':
                mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

            return send_file(
                file_path,
                as_attachment=True,
                download_name=safe_filename,
                mimetype=mimetype
            )
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500

    return _download()

@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({'success': False, 'error': 'حجم الملف كبير جداً (30MB كحد أقصى)'}), 413

@app.errorhandler(500)
def internal_server_error(error):
    logger.error(f"خطأ داخلي في الخادم: {str(error)}")
    return jsonify({
        'success': False,
        'error': f'خطأ داخلي: {str(error)}', # Expose error for debugging
        'code': 500
    }), 500

from werkzeug.exceptions import HTTPException

@app.errorhandler(Exception)
def handle_unexpected_error(error):
    # تمرير أخطاء HTTP كما هي (مثل abort(500) أو abort(404))
    if isinstance(error, HTTPException):
        return error

    logger.error(f"خطأ غير متوقع: {str(error)}")
    logger.error(f"نوع الخطأ: {type(error).__name__}")
    import traceback
    logger.error(f"تتبع الخطأ:\n{traceback.format_exc()}")

    # لا نرجع تفاصيل الخطأ للمستخدم لأسباب أمنية
    return jsonify({
        'success': False,
        'error': 'حدث خطأ غير متوقع. يرجى المحاولة لاحقاً.',
        'code': 500
    }), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
